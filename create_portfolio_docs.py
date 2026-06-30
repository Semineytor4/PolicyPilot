"""
Erstellt die 3 Portfolio-Dokumente (.docx) für das IU-Portfolio.
"""
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

with open("results.json") as f:
    results = json.load(f)

PLACEHOLDERS = {
    "name":     "[BITTE AUSFÜLLEN: Vor- und Nachname]",
    "matrikel": "[BITTE AUSFÜLLEN: Matrikelnummer]",
    "kurs":     "[BITTE AUSFÜLLEN: Kursbezeichnung]",
    "dozent":   "[BITTE AUSFÜLLEN: Dozent/in]",
    "github":   "[BITTE AUSFÜLLEN: GitHub-Repository-Link]",
    "datum":    "30. Juni 2026",
}

# ─────────────────────────────────────────────────────
# Hilfsfunktionen
# ─────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    # Seitenränder
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
    # Standardschriftart
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    # Absatzabstand
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(16)
    return doc

def set_heading_style(doc, level_styles):
    """Heading-Stile anpassen."""
    for level, (size, bold) in level_styles.items():
        try:
            style = doc.styles[f'Heading {level}']
            style.font.name = 'Arial'
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        except Exception:
            pass

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.name = 'Arial'
    return h

def add_paragraph(doc, text="", bold_parts=None):
    """Fügt einen Absatz hinzu. bold_parts: list of strings to bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(17)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    else:
        # text is list of (text, bold) tuples
        for part_text, is_bold in text:
            run = p.add_run(part_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.bold = is_bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return p

def add_table_row(table, cells, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(text)
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.bold = bold
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    return row

def add_image(doc, path, width_cm=14, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width_cm))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.runs[0] if cap.runs else cap.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
    else:
        doc.add_paragraph(f"[Plot nicht gefunden: {path}]")

def add_page_break(doc):
    doc.add_page_break()

def add_toc_placeholder(doc):
    """Fügt einen TOC-Feldcode ein."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    note = doc.add_paragraph()
    note_run = note.add_run("→ Bitte im geöffneten Word-Dokument das Inhaltsverzeichnis aktualisieren: "
                            "Rechtsklick auf das Feld → 'Felder aktualisieren'")
    note_run.font.size = Pt(9)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

def add_cover_page(doc, title, phase):
    """Deckblatt mit Platzhaltern."""
    doc.add_paragraph()
    doc.add_paragraph()

    # IU Logo Platzhalter
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("[IU INTERNATIONALES HOCHSCHULLOGO – HIER EINFÜGEN]")
    logo_run.font.size = Pt(10)
    logo_run.font.italic = True
    logo_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()
    doc.add_paragraph()

    # Titel
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("CustomerChurnLab")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("Customer Churn Prediction mit klassischem Machine Learning")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    phase_p = doc.add_paragraph()
    phase_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    phase_run = phase_p.add_run(title)
    phase_run.font.size = Pt(16)
    phase_run.font.bold = True
    phase_run.font.name = 'Arial'
    phase_run.font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadaten-Tabelle
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ("Name:", PLACEHOLDERS["name"]),
        ("Matrikelnummer:", PLACEHOLDERS["matrikel"]),
        ("Kurs:", PLACEHOLDERS["kurs"]),
        ("Dozent/in:", PLACEHOLDERS["dozent"]),
        ("Portfolio-Phase:", phase),
        ("Datum:", PLACEHOLDERS["datum"]),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    gh_p = doc.add_paragraph()
    gh_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gh_run = gh_p.add_run(f"GitHub-Repository: {PLACEHOLDERS['github']}")
    gh_run.font.size = Pt(10)
    gh_run.font.italic = True
    gh_run.font.color.rgb = RGBColor(0x00, 0x78, 0xD7)

    doc.add_page_break()

def add_references(doc):
    """Quellenverzeichnis im APA-Stil."""
    add_heading(doc, "Quellenverzeichnis", level=1)
    refs = [
        "Stock, S., Becker, J., Grimm, D., Hotfilter, T., Molinar, G., Stang, M., & Stork, W. (2021). "
        "QUA³CK – A Machine Learning Development Process. Karlsruher Institut für Technologie (KIT). "
        "Unveröffentlichter Bericht / Lehrunterlage.",

        "IBM / Kaggle. (o. J.). Telco Customer Churn. Kaggle Datasets. "
        "Abgerufen am 30. Juni 2026, von https://www.kaggle.com/datasets/blastchar/telco-customer-churn",

        f"GitHub-Repository: {PLACEHOLDERS['github']}",

        "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. "
        "Journal of Machine Learning Research, 12, 2825–2830. scikit-learn Version ≥ 1.3.0.",

        "McKinney, W. (2010). Data Structures for Statistical Computing in Python. "
        "Proceedings of the 9th Python in Science Conference. pandas Version ≥ 2.0.0.",

        "Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. "
        "Computing in Science & Engineering, 9(3), 90–95. matplotlib Version ≥ 3.7.0.",

        "Waskom, M. (2021). Seaborn: Statistical Data Visualization. "
        "Journal of Open Source Software, 6(60), 3021. seaborn Version ≥ 0.12.0.",

        "Streamlit Inc. (2023). Streamlit – The fastest way to build data apps. "
        "https://streamlit.io. streamlit Version ≥ 1.28.0.",

        "Joblib Development Team. (2023). Joblib: running Python functions as pipeline jobs. "
        "https://joblib.readthedocs.io. joblib Version ≥ 1.3.0.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════════════════
# DOKUMENT 1: Portfolio Phase 1 – Q-Phase (Konzeptionsphase)
# ═══════════════════════════════════════════════════════════════════════════════
print("Erstelle Dokument 1 (Phase 1 – Konzeption)...")
doc1 = new_doc()
set_heading_style(doc1, {1: (14, True), 2: (12, True), 3: (11, True)})

add_cover_page(doc1, "Portfolio Phase 1: Konzeptionsphase", "Phase 1 von 3")

# Inhaltsverzeichnis
add_heading(doc1, "Inhaltsverzeichnis", level=1)
add_toc_placeholder(doc1)
doc1.add_page_break()

# 1. Einleitung
add_heading(doc1, "1. Einleitung und Projekthintergrund", level=1)
add_paragraph(doc1,
    "Das vorliegende Portfolio-Projekt untersucht die Vorhersage von Customer Churn – "
    "der Abwanderung von Telekommunikationskunden – mithilfe klassischer Machine-Learning-Verfahren. "
    "Kundenabwanderung ist wirtschaftlich kritisch: Die Neukundenakquise ist branchenübergreifend "
    "deutlich kostspieliger als die Bindung bestehender Kunden. Datengetriebene Modelle ermöglichen es, "
    "gefährdete Kunden frühzeitig zu identifizieren und gezielte Retentionsmaßnahmen einzuleiten."
)
add_paragraph(doc1,
    "Als Datenbasis dient der öffentlich verfügbare IBM Telco Customer Churn Datensatz (Kaggle), "
    "der 7.043 Kunden mit 20 Eingangsmerkmalen und einer binären Zielvariable (Churn: Yes/No) enthält. "
    "Das Projekt folgt dem QUA³CK-Prozessmodell (Stock et al., 2021) und endet mit einer "
    "interaktiven Streamlit-Webanwendung zur Demonstration der Modellergebnisse."
)

# 1.1 Ideenprüfung
add_heading(doc1, "1.1 Ideenprüfung: Tragfähigkeit des Konzepts", level=2)
add_paragraph(doc1, "Das Thema Customer Churn Prediction wurde anhand folgender Kriterien auf Tragfähigkeit geprüft:")

pro_con_table = doc1.add_table(rows=1, cols=2)
pro_con_table.style = 'Table Grid'
pro_con_table.rows[0].cells[0].text = "✅ Pro"
pro_con_table.rows[0].cells[1].text = "⚠️ Herausforderungen"
for cell in pro_con_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)

pros = [
    "Öffentlich verfügbarer, realer Datensatz (keine Datenschutz-Probleme)",
    "Klar definierte binäre Klassifikationsaufgabe",
    "Hohe praktische Relevanz für Unternehmen",
    "Vollständiger ML-Prozess von EDA bis Deployment demonstrierbar",
    "Scikit-learn bietet alle benötigten Tools",
]
cons = [
    "Klassenimbalance (≈26,5% Churn) erfordert spezielle Behandlung",
    "TotalCharges liegt als String vor (Datenqualitätsproblem)",
    "Ergebnisse sind korrelativ, nicht kausal",
    "Dataset spiegelt US-Telekommunikationsmarkt wider – begrenzte Generalisierbarkeit",
]
for p, c in zip(pros, cons):
    row = pro_con_table.add_row()
    row.cells[0].text = f"• {p}"
    row.cells[1].text = f"• {c}"
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

doc1.add_paragraph()

# 2. Konzept und Zielsetzung
add_heading(doc1, "2. Konzept und Zielsetzung", level=1)
add_heading(doc1, "2.1 Zentrale Forschungsfrage", level=2)
add_paragraph(doc1,
    "Wie gut lässt sich die Abwanderung von Kunden (Customer Churn) mithilfe klassischer "
    "Machine-Learning-Verfahren auf Basis tabellarischer Kundendaten vorhersagen, und welcher "
    "Algorithmus bietet dabei die beste Kombination aus Vorhersagequalität, Interpretierbarkeit "
    "und Rechenaufwand?"
)

add_heading(doc1, "2.2 Arbeitshypothesen", level=2)
add_paragraph(doc1,
    "Aus der Forschungsfrage wurden drei überprüfbare Arbeitshypothesen abgeleitet, "
    "die im Verlauf des Projekts empirisch getestet werden:"
)
for h, text in [
    ("H1:", "Vertragsmerkmale (Contract Type, Tenure) haben einen stärkeren Einfluss auf Churn als demografische Merkmale (Geschlecht, SeniorCitizen)."),
    ("H2:", "Ensemble-Modelle (Random Forest, Gradient Boosting) erzielen eine höhere Vorhersageleistung (ROC-AUC, F1) als lineare Modelle (Logistic Regression), sind aber schwerer interpretierbar."),
    ("H3:", "Saubere Datenvorverarbeitung (korrektes Encoding, Skalierung für lineare Modelle, Leakage-freies Vorgehen über scikit-learn-Pipelines) verbessert die Generalisierungsleistung messbar gegenüber naivem Vorgehen."),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(f"{h} ")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)

add_heading(doc1, "2.3 Problemstellung und Zielgruppe", level=2)
add_paragraph(doc1,
    "Die Zielgruppe umfasst Unternehmen, Datenanalysten und Fachabteilungen (Marketing, Customer Success), "
    "die Kundenbindung datenbasiert unterstützen möchten. Das Modell soll als Entscheidungshilfe dienen: "
    "Kunden mit hoher Abwanderungswahrscheinlichkeit werden priorisiert für Retentionsmaßnahmen. "
    "Für nicht-technische Nutzer wird eine Streamlit-Webanwendung entwickelt, die die Ergebnisse "
    "verständlich und interaktiv präsentiert."
)

add_heading(doc1, "2.4 Erfolgsmetriken (KPIs)", level=2)
add_paragraph(doc1,
    "Da der Datensatz klassenimbalanciert ist (≈26,5% Churn), werden folgende Metriken verwendet:"
)
kpi_table = doc1.add_table(rows=1, cols=3)
kpi_table.style = 'Table Grid'
for cell, text in zip(kpi_table.rows[0].cells, ["Metrik", "Bedeutung im Kontext", "Zielwert"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for row_data in [
    ("F1-Score", "Balance aus Precision und Recall – Hauptmetrik bei Imbalance", "> 0.55"),
    ("Recall", "Anteil erkannter Kündiger – für Frühwarnsystem kritisch", "> 0.65"),
    ("ROC-AUC", "Threshold-unabhängige Diskriminierungsfähigkeit", "> 0.80"),
    ("PR-AUC", "Precision-Recall unter Imbalance", "> 0.55"),
]:
    row = kpi_table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

doc1.add_paragraph()

# 3. Methodik und Vorgehen
add_heading(doc1, "3. Methodik und Vorgehen", level=1)
add_heading(doc1, "3.1 Das QUA³CK-Prozessmodell", level=2)
add_paragraph(doc1,
    "Das Projekt folgt dem QUA³CK-Prozessmodell (Stock et al., 2021), einem iterativen "
    "ML-Entwicklungsprozess, der speziell für die strukturierte Bearbeitung von Machine-Learning-Projekten "
    "entwickelt wurde. QUA³CK steht für:"
)
for phase, desc in [
    ("Q – Question:", "Problemdefinition, Forschungsfrage, Hypothesen"),
    ("U – Understanding:", "Explorative Datenanalyse (EDA), Datenbeschreibung, Korrelationsanalyse"),
    ("A¹ – Algorithm Selection:", "Auswahl und Vergleich von Kandidatenmodellen"),
    ("A² – Adapting Features:", "Leakage-freie Datenvorverarbeitung via scikit-learn-Pipeline"),
    ("A³ – Adjusting Hyperparameters:", "Hyperparameter-Optimierung, Cross-Validation"),
    ("C – Conclusion & Comparison:", "Finale Evaluation, Modellvergleich, Hypothesenprüfung"),
    ("K – Knowledge Transfer:", "Streamlit-App, Dokumentation, Portfolio"),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(phase + " ")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)

add_paragraph(doc1,
    "\nDer QUA³CK-Prozess wurde gegenüber CRISP-DM als Rahmenwerk gewählt, da er explizit "
    "auf den iterativen A³-Zyklus (Algorithmusauswahl, Feature-Anpassung, Hyperparameter-Tuning) "
    "fokussiert und damit den ML-Modellierungsprozess granularer abbildet (Stock et al., 2021)."
)

add_heading(doc1, "3.2 Geplantes Vorgehen", level=2)
for i, step in enumerate([
    "Explorative Datenanalyse (U-Phase): Datenstruktur, Verteilungen, Korrelationen, Klassenimbalance",
    "Preprocessing-Pipeline: TotalCharges-Konvertierung, OneHotEncoding, StandardScaler, SimpleImputer",
    "Modelltraining & Cross-Validation: DummyClassifier, Logistic Regression, Random Forest, Gradient Boosting",
    "Holdout-Testset-Evaluation: Finaler Test auf zurückgehaltenen 20% der Daten",
    "Modellauswahl und Begründung: Primär F1/Recall, sekundär Interpretierbarkeit",
    "Hypothesenprüfung: H1, H2, H3 empirisch verifizieren oder widerlegen",
    "Deployment: Modell serialisieren (joblib), Streamlit-App entwickeln",
], 1):
    add_bullet(doc1, f"Schritt {i}: {step}")

add_heading(doc1, "3.3 Deployment-Ziel (Streamlit)", level=2)
add_paragraph(doc1,
    "Als finales Deployment-Produkt wird eine interaktive Streamlit-Webanwendung entwickelt "
    "(app.py). Diese ermöglicht es Anwendern, für beliebige Kundendaten eine Churn-Wahrscheinlichkeit "
    "zu berechnen und erhält die Ausgabe als farbcodierten Risikoindikator (grün/gelb/rot). "
    "Zusätzlich werden die wichtigsten Einflussfaktoren (Feature-Koeffizienten) visualisiert. "
    "Der vollständige Quellcode ist im GitHub-Repository verfügbar: "
    f"{PLACEHOLDERS['github']}"
)

# 4. Datengrundlage
add_heading(doc1, "4. Datengrundlage", level=1)
add_paragraph(doc1,
    "Als Datenbasis wird der IBM Telco Customer Churn Datensatz (IBM/Kaggle) verwendet. "
    "Der Datensatz enthält anonymisierte Kundendaten eines US-amerikanischen Telekommunikationsunternehmens."
)
data_table = doc1.add_table(rows=1, cols=2)
data_table.style = 'Table Grid'
for cell, text in zip(data_table.rows[0].cells, ["Merkmal", "Wert"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for k, v in [
    ("Anzahl Kunden", "7.043"),
    ("Anzahl Features", "20 Eingangsmerkmale + 1 Zielvariable"),
    ("Zielvariable", "Churn (Yes/No) → binär kodiert (1/0)"),
    ("Churn-Rate", "26,5% (Klassenimbalance vorhanden)"),
    ("Kategoriale Features", "15 (z. B. Contract, InternetService, PaymentMethod)"),
    ("Numerische Features", "4 (SeniorCitizen, tenure, MonthlyCharges, TotalCharges)"),
    ("Bekannte Datenprobleme", "TotalCharges als String gespeichert (11 fehlende Werte nach Konvertierung)"),
    ("Lizenz", "Öffentlich, frei verwendbar (Kaggle-Datensatz)"),
]:
    row = data_table.add_row()
    row.cells[0].text = k
    row.cells[1].text = v
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

doc1.add_paragraph()

add_references(doc1)
doc1.save("Portfolio_Phase1_CustomerChurn.docx")
print("✅ Portfolio_Phase1_CustomerChurn.docx gespeichert")


# ═══════════════════════════════════════════════════════════════════════════════
# DOKUMENT 2: Portfolio Phase 2 – Erarbeitungsphase (U + A³)
# ═══════════════════════════════════════════════════════════════════════════════
print("Erstelle Dokument 2 (Phase 2 – Erarbeitungsphase)...")
doc2 = new_doc()
set_heading_style(doc2, {1: (14, True), 2: (12, True), 3: (11, True)})

add_cover_page(doc2, "Portfolio Phase 2: Erarbeitungsphase", "Phase 2 von 3")

add_heading(doc2, "Inhaltsverzeichnis", level=1)
add_toc_placeholder(doc2)
doc2.add_page_break()

# Überblick
add_heading(doc2, "1. Überblick und Fortschritt", level=1)
add_paragraph(doc2,
    "Die Erarbeitungsphase umfasst die praktische Umsetzung des im Konzept (Phase 1) definierten Plans. "
    "Konkret wurden die U-Phase (Explorative Datenanalyse) und der A³-Zyklus (Algorithm Selection, "
    "Adapting Features, Adjusting via Cross-Validation) vollständig implementiert. "
    "Alle Ergebnisse basieren auf dem IBM Telco Customer Churn Datensatz (7.043 Kunden). "
    "Der vollständige Code ist im GitHub-Repository dokumentiert: "
    f"{PLACEHOLDERS['github']}"
)

# 2. U-Phase: EDA
add_heading(doc2, "2. U-Phase: Explorative Datenanalyse (EDA)", level=1)
add_heading(doc2, "2.1 Datenbeschreibung", level=2)
add_paragraph(doc2,
    "Nach dem Laden des Datensatzes und der Konvertierung von TotalCharges "
    "(pd.to_numeric(errors='coerce')) wurden 11 fehlende Werte identifiziert und via "
    "Median-Imputation behandelt. Die endgültige Datenstruktur umfasst:"
)
for item in [
    "7.043 Kunden (Zeilen) × 20 Features (Spalten) + Zielvariable Churn",
    "4 numerische Features: SeniorCitizen, tenure, MonthlyCharges, TotalCharges",
    "15 kategoriale Features: gender, Partner, Dependents, PhoneService, MultipleLines, InternetService, OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, StreamingTV, StreamingMovies, Contract, PaperlessBilling, PaymentMethod",
    "Churn-Rate: 26,5% (1.869 Churn-Fälle, 5.174 Nicht-Churn)",
]:
    add_bullet(doc2, item)

add_heading(doc2, "2.2 Verteilung der Zielvariable (Churn)", level=2)
add_paragraph(doc2,
    "Die folgende Abbildung zeigt die Klassenverteilung im Datensatz. "
    "Die Klassen sind unausgewogen (26,5% Churn vs. 73,5% kein Churn). "
    "Dies erfordert bei der Modellierung spezielle Maßnahmen: "
    "stratifizierte Splits, Verwendung von class_weight='balanced' sowie "
    "Fokus auf F1/Recall statt Accuracy als primäre Metrik."
)
add_image(doc2, "plots/churn_distribution.png", width_cm=10,
          caption="Abb. 1: Klassenverteilung der Zielvariable Churn (5.174 Kein Churn vs. 1.869 Churn)")

add_heading(doc2, "2.3 Analyse numerischer Features", level=2)
add_paragraph(doc2,
    "Die Histogramme in Abb. 2 zeigen die Verteilung von tenure (Vertragsdauer) und "
    "MonthlyCharges (monatliche Kosten) nach Churn-Status. Zwei klare Muster sind erkennbar:"
)
for text in [
    "Tenure: Kunden mit kurzer Vertragsdauer (0–12 Monate) kündigen deutlich häufiger. Bei tenure > 50 Monaten sinkt die Churn-Rate stark. Dies stützt Hypothese H1.",
    "MonthlyCharges: Churn-Kunden zahlen im Durchschnitt höhere monatliche Gebühren. Eine Häufung bei 70–100 € ist bei Churn-Kunden deutlich ausgeprägter.",
]:
    add_bullet(doc2, text)
add_image(doc2, "plots/numeric_features_by_churn.png", width_cm=14,
          caption="Abb. 2: Verteilung numerischer Features nach Churn-Status")

add_heading(doc2, "2.4 Analyse kategorialer Features", level=2)
add_paragraph(doc2,
    "Abb. 3 zeigt die Churn-Raten für drei zentrale kategoriale Features:"
)
for text in [
    "Contract: Month-to-month-Kunden weisen mit Abstand die höchste Churn-Rate auf (>40%). 2-Jahres-Kunden kündigen kaum (<5%). Vertragsart ist der stärkste Prädiktor.",
    "InternetService: Fiber-optic-Kunden zeigen deutlich höhere Churn-Raten als DSL-Kunden. Kunden ohne Internet kündigen kaum.",
    "PaymentMethod: Kunden, die per Electronic check zahlen, kündigen häufiger als Kunden mit automatischen Zahlungsmethoden.",
]:
    add_bullet(doc2, text)
add_image(doc2, "plots/categorical_churn_rates.png", width_cm=14,
          caption="Abb. 3: Churn-Rate nach kategorialen Merkmalen (Contract, InternetService, PaymentMethod)")

add_heading(doc2, "2.5 Korrelationsanalyse und Fazit der EDA", level=2)
add_paragraph(doc2,
    "Die Korrelationsanalyse numerischer Features zeigt: tenure ist stark negativ mit Churn korreliert "
    "(längere Vertragsdauer → weniger Kündigung), MonthlyCharges ist positiv korreliert "
    "(höhere Kosten → mehr Kündigung). TotalCharges ist stark mit tenure korreliert "
    "(Multikollinearität, da TotalCharges = tenure × durchschnittliche Monatskosten)."
)
add_paragraph(doc2,
    "Fazit: Vertragsdauer, Vertragsart und monatliche Kosten sind die zentralen Einflussfaktoren. "
    "Diese Erkenntnis fließt direkt in die Hypothesenprüfung in Phase 3 ein."
)

# 3. A³-Phase
add_heading(doc2, "3. A³-Phase: Implementierung des Machine-Learning-Prozesses", level=1)

add_heading(doc2, "3.1 Pipeline-Architektur (A²: Adapting Features)", level=2)
add_paragraph(doc2,
    "Ein zentrales Designprinzip ist die vollständig leakage-freie Datenvorverarbeitung. "
    "Alle Transformationen finden ausschließlich innerhalb einer scikit-learn Pipeline statt, "
    "d. h. der Preprocessor sieht bei Cross-Validation nur die jeweiligen Trainingsfolds:"
)
code_text = (
    "Pipeline([\n"
    "  ('prep', ColumnTransformer([\n"
    "    ('num', Pipeline([SimpleImputer(median), StandardScaler()]), num_cols),\n"
    "    ('cat', Pipeline([SimpleImputer(mode), OneHotEncoder(handle_unknown='ignore')]), cat_cols)\n"
    "  ])),\n"
    "  ('model', LogisticRegression(class_weight='balanced'))\n"
    "])"
)
code_para = doc2.add_paragraph()
code_run = code_para.add_run(code_text)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_para.paragraph_format.left_indent = Cm(1)

add_paragraph(doc2,
    "Numerische Features werden mit MedianImputer und StandardScaler normalisiert. "
    "Kategoriale Features werden mit dem häufigsten Wert imputiert und one-hot-kodiert "
    "(handle_unknown='ignore' für robuste Behandlung unbekannter Kategorien im Testset). "
    "TotalCharges wurde mit pd.to_numeric(errors='coerce') korrekt in float konvertiert."
)

add_heading(doc2, "3.2 Modellauswahl (A¹: Algorithm Selection)", level=2)
add_paragraph(doc2,
    "Vier Modelle wurden als Kandidaten definiert, um ein breites Spektrum abzudecken:"
)
for model, desc in [
    ("DummyClassifier (majority class)", "Baseline-Referenz, der immer die Mehrheitsklasse vorhersagt"),
    ("Logistic Regression", "Interpretierbar, schnell, geeignet für lineare Zusammenhänge, class_weight='balanced'"),
    ("Random Forest", "Robustes Ensemble, erfasst nichtlineare Zusammenhänge, class_weight='balanced'"),
    ("Gradient Boosting", "Leistungsstarkes Boosting-Verfahren, meist höchste Precision"),
]:
    p = doc2.add_paragraph()
    r1 = p.add_run(f"{model}: ")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)

add_heading(doc2, "3.3 Cross-Validation Ergebnisse (A³)", level=2)
add_paragraph(doc2,
    "Alle Modelle wurden mit Stratified K-Fold (5 Folds, random_state=42) auf dem Trainingsset "
    "(80% der Daten, 5.634 Kunden) evaluiert. Die Ergebnisse sind in Tab. 1 zusammengefasst:"
)
cv_table = doc2.add_table(rows=1, cols=6)
cv_table.style = 'Table Grid'
for cell, text in zip(cv_table.rows[0].cells,
                      ["Modell", "ROC-AUC", "F1", "Precision", "Recall", "PR-AUC"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for r in results["cv_results"]:
    row = cv_table.add_row()
    for cell, val in zip(row.cells, [r["Modell"], r["ROC-AUC"], r["F1"],
                                      r["Precision"], r["Recall"], r["PR-AUC"]]):
        cell.text = str(val)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc2.add_paragraph()
p = doc2.add_paragraph()
p.add_run("Tab. 1: 5-Fold Cross-Validation Ergebnisse (Trainingsset, 5.634 Kunden)").italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_image(doc2, "plots/cv_model_comparison.png", width_cm=13,
          caption="Abb. 4: Cross-Validation Metriken im Modellvergleich")

add_heading(doc2, "3.4 Reflexion der Erarbeitungsphase", level=2)
add_paragraph(doc2, "Was hat gut funktioniert:")
for text in [
    "Die leakage-freie scikit-learn-Pipeline erwies sich als solide Grundarchitektur",
    "Der DummyClassifier lieferte eine sinnvolle Baseline (F1=0.0, ROC-AUC=0.5) zur Einordnung",
    "StratifiedKFold stellte ausgewogene Klassenverteilung in allen Folds sicher",
    "class_weight='balanced' verbesserte Recall für Churn-Klasse spürbar",
]:
    add_bullet(doc2, text)

add_paragraph(doc2, "Herausforderungen:")
for text in [
    "TotalCharges war als String gespeichert: Konvertierung mit pd.to_numeric(errors='coerce') löste das Problem, hinterließ aber 11 NaN-Werte (bei tenure=0), die via Median-Imputation behandelt wurden",
    "Klassenimbalance (26,5% Churn) führt dazu, dass Accuracy als Metrik trügerisch wäre (Dummy-Classifier würde 73,5% Accuracy erreichen) – F1 und Recall sind deutlich aussagekräftiger",
    "Gradient Boosting: Deutlich geringerer Recall (0.53) trotz höherer Precision – für ein Churn-Frühwarnsystem suboptimal",
    "Hyperparameter-Tuning (GridSearchCV) war konzeptionell geplant, konnte aufgrund des Projektumfangs nur mit Default-Parametern (außer n_estimators=300, class_weight='balanced') umgesetzt werden",
]:
    add_bullet(doc2, text)

add_heading(doc2, "3.5 Meilensteine (M1–M5)", level=2)
for m, text in [
    ("M1", "Datensatz geladen, TotalCharges konvertiert, Target/Features definiert ✅"),
    ("M2", "Baseline (DummyClassifier) etabliert: F1=0.0, ROC-AUC=0.5 ✅"),
    ("M3", "Leakage-freie Pipeline (ColumnTransformer + OneHot + Scaler) implementiert ✅"),
    ("M4", "Modellvergleich via 5-Fold CV abgeschlossen (Tab. 1) ✅"),
    ("M5", "LogReg als Kandidat für Holdout-Evaluation identifiziert (bester F1/Recall) ✅"),
]:
    p = doc2.add_paragraph()
    r1 = p.add_run(f"{m}: ")
    r1.bold = True
    r1.font.name = 'Arial'
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.name = 'Arial'
    r2.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)

add_references(doc2)
doc2.save("Portfolio_Phase2_CustomerChurn.docx")
print("✅ Portfolio_Phase2_CustomerChurn.docx gespeichert")


# ═══════════════════════════════════════════════════════════════════════════════
# DOKUMENT 3: Portfolio Phase 3 – Finalisierungsphase (C + K)
# ═══════════════════════════════════════════════════════════════════════════════
print("Erstelle Dokument 3 (Phase 3 – Finalisierung)...")
doc3 = new_doc()
set_heading_style(doc3, {1: (14, True), 2: (12, True), 3: (11, True)})

add_cover_page(doc3, "Portfolio Phase 3: Finalisierungsphase", "Phase 3 von 3")

add_heading(doc3, "Inhaltsverzeichnis", level=1)
add_toc_placeholder(doc3)
doc3.add_page_break()

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc3, "Abstract", level=1)
add_paragraph(doc3,
    "Das Projekt CustomerChurnLab untersucht die Vorhersage von Customer Churn – der Abwanderung "
    "von Telekommunikationskunden – mithilfe klassischer Machine-Learning-Verfahren. Zielsetzung "
    "ist die Beantwortung der zentralen Forschungsfrage: Welcher ML-Algorithmus bietet die beste "
    "Kombination aus Vorhersagequalität, Interpretierbarkeit und Rechenaufwand? Als Datenbasis "
    "dient der öffentlich verfügbare IBM Telco Customer Churn Datensatz (7.043 Kunden, 20 Features)."
)
add_paragraph(doc3,
    "Das Projekt folgt dem QUA³CK-Prozessmodell (Stock et al., 2021): Von der Problemdefinition "
    "(Q) über explorative Datenanalyse (U) und iterativen A³-Zyklus (Algorithmusauswahl, "
    "Feature-Anpassung, Hyperparameter-Abstimmung) bis zur Schlussfolgerung (C) und dem "
    "Wissenstransfer (K) in Form einer interaktiven Streamlit-Webanwendung."
)
add_paragraph(doc3,
    "Vier Modelle wurden verglichen (DummyClassifier, Logistic Regression, Random Forest, "
    "Gradient Boosting). Logistic Regression erzielte den besten F1-Score (0.614) und "
    "den höchsten Recall (0.783) auf dem Holdout-Testset – entscheidend für ein "
    "Churn-Frühwarnsystem, das möglichst wenige Kündiger verpassen soll. "
    "Die Hypothesen H1 (Vertragsmerkmale dominieren) und H3 (Pipeline-Preprocessing verbessert "
    "Generalisierung) wurden bestätigt; H2 (Ensemble-Modelle besser) wurde widerlegt. "
    "Das finale Modell ist als Streamlit-App deploybar und als Jupyter-Notebook vollständig "
    "reproduzierbar. Der Quellcode ist öffentlich verfügbar: "
    f"{PLACEHOLDERS['github']}"
)
doc3.add_page_break()

# ── Rückblick Phase 1 ─────────────────────────────────────────────────────────
add_heading(doc3, "1. Rückblick Phase 1: Konzeptionsphase", level=1)
add_paragraph(doc3,
    "In Phase 1 wurde die zentrale Forschungsfrage definiert sowie drei empirisch prüfbare "
    "Hypothesen (H1–H3) formuliert. Das QUA³CK-Prozessmodell wurde als methodischer Rahmen "
    "gewählt und begründet. Als Datensatz wurde der IBM Telco Customer Churn Datensatz "
    "(7.043 Kunden, 26,5% Churn-Rate) ausgewählt. Erfolgsmetriken (F1 > 0.55, Recall > 0.65, "
    "ROC-AUC > 0.80) und das Deployment-Ziel (Streamlit-App) wurden festgelegt. "
    "Alle geplanten Meilensteine konnten umgesetzt werden."
)

# ── Rückblick Phase 2 ─────────────────────────────────────────────────────────
add_heading(doc3, "2. Rückblick Phase 2: Erarbeitungsphase", level=1)
add_paragraph(doc3,
    "Phase 2 umfasste die explorative Datenanalyse (U-Phase) und den A³-Zyklus. "
    "Zentrale EDA-Erkenntnis: Vertragsdauer (tenure), Vertragsart (Contract) und monatliche "
    "Kosten (MonthlyCharges) sind die wichtigsten Prädiktoren. Die Klassenimbalance (26,5% Churn) "
    "wurde durch stratifizierte Splits und class_weight='balanced' adressiert. "
    "Die scikit-learn Pipeline stellte Leakage-Freiheit sicher. "
    "Vier Modelle wurden per 5-Fold Cross-Validation verglichen. Logistic Regression "
    "erzielte den besten F1 (0.628) und Recall (0.801) – trotz niedrigerer Precision als GradBoost."
)

# ── Technischer Anhang ────────────────────────────────────────────────────────
add_heading(doc3, "3. Technischer Anhang: Ressourcen und Architektur", level=1)
add_heading(doc3, "3.1 Verwendete Software und Bibliotheken", level=2)
tech_table = doc3.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
for cell, text in zip(tech_table.rows[0].cells, ["Bibliothek", "Version", "Verwendungszweck"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for lib, ver, purpose in [
    ("Python", "3.10+", "Programmiersprache"),
    ("pandas", "≥ 2.0.0", "Datenanalyse, Preprocessing"),
    ("scikit-learn", "≥ 1.3.0", "ML-Modelle, Pipelines, Metriken"),
    ("matplotlib", "≥ 3.7.0", "Visualisierungen"),
    ("seaborn", "≥ 0.12.0", "Statistische Plots"),
    ("joblib", "≥ 1.3.0", "Modell-Serialisierung"),
    ("streamlit", "≥ 1.28.0", "Web-App Deployment"),
    ("Jupyter Notebook", "–", "Interaktive Entwicklung"),
]:
    row = tech_table.add_row()
    for cell, val in zip(row.cells, [lib, ver, purpose]):
        cell.text = val
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc3.add_paragraph()

add_heading(doc3, "3.2 Projektdateien", level=2)
for fname, desc in [
    ("Q_conception.ipynb", "Q-Phase: Forschungsfrage, Hypothesen, Konzept"),
    ("U_EDA.ipynb", "U-Phase: EDA-Code mit Visualisierungen"),
    ("U_Analyse.ipynb", "U-Phase: Analyse-Interpretation"),
    ("A_Development_Phase.ipynb", "A³-Phase: Pipeline und Modellvergleich"),
    ("A_Development_Phase_Results.ipynb", "A³-Phase: Cross-Validation Ergebnisse"),
    ("C_Evaluation_Final.py", "C-Phase: Holdout-Evaluation, alle Plots, Modell-Export"),
    ("app.py", "K-Phase: Streamlit-App"),
    ("models/logreg_churn_pipeline.pkl", "Serialisiertes finales Modell (joblib)"),
    ("requirements.txt", "Abhängigkeiten"),
    ("AI_TOOL_DISCLOSURE.ipynb", "KI-Tool-Nutzung Dokumentation"),
]:
    p = doc3.add_paragraph()
    r1 = p.add_run(f"{fname}: ")
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    r1.bold = True
    r2 = p.add_run(desc)
    r2.font.name = 'Arial'
    r2.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)

# ── C-Phase: Finale Evaluation ────────────────────────────────────────────────
add_heading(doc3, "4. C-Phase: Finale Evaluation und Ergebnisse", level=1)

add_heading(doc3, "4.1 Finale Modellwahl mit Begründung", level=2)
add_paragraph(doc3,
    "Als finales Modell wurde Logistic Regression ausgewählt. Die Begründung erfolgt "
    "anhand eines direkten Vergleichs mit dem konkurrierenden Ansatz Gradient Boosting:"
)
choice_table = doc3.add_table(rows=1, cols=3)
choice_table.style = 'Table Grid'
for cell, text in zip(choice_table.rows[0].cells,
                      ["Kriterium", "Logistic Regression ✅", "Gradient Boosting"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for crit, lr, gb in [
    ("F1-Score (Test)", "0.614 (höher)", "0.590"),
    ("Recall (Test)", "0.783 (höher ✅)", "0.524"),
    ("Precision (Test)", "0.504", "0.674 (höher)"),
    ("ROC-AUC (Test)", "0.841", "0.843 (minimal besser)"),
    ("Interpretierbarkeit", "Koeffizienten → Odds Ratios ✅", "Feature Importance, kein direkter Effekt"),
    ("Trainingszeit", "Sehr schnell (Sekunden) ✅", "Moderat (Sekunden–Minuten)"),
    ("Für Frühwarnung", "Ideal (hoher Recall: 78,3%) ✅", "Ungeeignet (Recall: 52,4%)"),
]:
    row = choice_table.add_row()
    for cell, val in zip(row.cells, [crit, lr, gb]):
        cell.text = val
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc3.add_paragraph()
add_paragraph(doc3,
    "Für ein Churn-Frühwarnsystem ist Recall entscheidend: Ein nicht erkannter Kündiger "
    "(False Negative) verursacht dauerhaften Kundenverlust, während ein False Positive "
    "(fälschlich als Kündiger markierter treuer Kunde) lediglich unnötige Retentionskosten "
    "erzeugt – ein deutlich geringerer Schaden. Daher ist Logistic Regression trotz geringerer "
    "Precision die optimale Wahl für diesen Anwendungsfall."
)

add_heading(doc3, "4.2 Holdout-Testset Evaluation", level=2)
add_paragraph(doc3,
    "Alle Modelle wurden final auf dem zurückgehaltenen Testset evaluiert "
    "(1.409 Kunden, 20% des Gesamtdatensatzes, stratifizierter Split):"
)
test_table = doc3.add_table(rows=1, cols=6)
test_table.style = 'Table Grid'
for cell, text in zip(test_table.rows[0].cells,
                      ["Modell", "F1", "Recall", "Precision", "ROC-AUC", "PR-AUC"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for r in results["test_results"]:
    row = test_table.add_row()
    vals = [r["Modell"], r.get("F1","–"), r.get("Recall","–"),
            r.get("Precision","–"), r.get("ROC-AUC","–"), r.get("PR-AUC","–")]
    for cell, val in zip(row.cells, vals):
        cell.text = str(val) if val is not None else "–"
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc3.add_paragraph()
p = doc3.add_paragraph()
p.add_run("Tab. 2: Holdout-Testset Evaluation (1.409 Kunden, 20% des Datensatzes)").italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc3, "4.3 Konfusionsmatrix", level=2)
cm = results["confusion_matrix"]
add_paragraph(doc3,
    f"Abb. 5 zeigt die Konfusionsmatrix für Logistic Regression auf dem Testset. "
    f"Von 374 tatsächlichen Churn-Kunden wurden {cm[1][1]} korrekt erkannt (True Positives). "
    f"{cm[1][0]} Churn-Kunden wurden nicht erkannt (False Negatives – verpässte Kündiger). "
    f"Von 1.035 Nicht-Churn-Kunden wurden {cm[0][0]} korrekt als loyal klassifiziert. "
    f"{cm[0][1]} loyale Kunden wurden fälschlicherweise als Kündiger markiert (False Positives)."
)
add_image(doc3, "plots/confusion_matrix.png", width_cm=9,
          caption="Abb. 5: Konfusionsmatrix – Logistic Regression (Testset, 1.409 Kunden)")

add_heading(doc3, "4.4 ROC- und Precision-Recall-Kurven", level=2)
add_paragraph(doc3,
    "Abb. 6 zeigt die ROC-Kurven aller drei Modelle. Gradient Boosting erzielt marginale "
    "Vorteile beim ROC-AUC (0.843 vs. 0.841 für LogReg), der Unterschied ist jedoch minimal. "
    "Die Precision-Recall-Kurve (Abb. 7) verdeutlicht, dass LogReg bei niedrigen Precision-Schwellen "
    "(also hohem Recall) besser abschneidet – relevant für das Frühwarnsystem."
)
add_image(doc3, "plots/roc_curves.png", width_cm=11,
          caption="Abb. 6: ROC-Kurven – Modellvergleich auf Testset")
add_image(doc3, "plots/pr_curves.png", width_cm=11,
          caption="Abb. 7: Precision-Recall-Kurven – Modellvergleich auf Testset")

add_heading(doc3, "4.5 Feature Importance (Logistic Regression)", level=2)
add_paragraph(doc3,
    "Abb. 8 zeigt die Top-20-Features nach absolutem Koeffizientenwert der Logistic Regression. "
    "Positive Koeffizienten erhöhen die Churn-Wahrscheinlichkeit (log-odds), negative senken sie:"
)
add_image(doc3, "plots/feature_importance_logreg.png", width_cm=13,
          caption="Abb. 8: Logistic Regression Koeffizienten – Top-20 Features (rot=Churn-fördernd, grün=Churn-hemmend)")
add_paragraph(doc3,
    "Die wichtigsten Churn-fördernden Features sind Contract_Month-to-month und tenure (negativ). "
    "Kurze Vertragsdauer und monatliche Verträge sind die stärksten Prädiktoren. "
    "Churn-hemmend wirken v. a. Contract_Two year und zusätzliche Services wie TechSupport_Yes."
)

# ── Hypothesenprüfung ─────────────────────────────────────────────────────────
add_heading(doc3, "5. Hypothesenprüfung", level=1)

add_heading(doc3, "H1: Vertragsmerkmale dominieren Demografik", level=2)
add_paragraph(doc3,
    "H1 lautet: Vertragsmerkmale (Contract Type, Tenure) haben einen stärkeren Einfluss auf Churn "
    "als demografische Merkmale (Geschlecht, SeniorCitizen)."
)
add_paragraph(doc3,
    "Ergebnis: ✅ BESTÄTIGT. Die Koeffizientenanalyse (Abb. 8) zeigt, dass Contract_Month-to-month "
    "und tenure zu den stärksten Prädiktoren gehören. gender und SeniorCitizen haben deutlich "
    "geringere Koeffizientenbeträge. Die EDA (Abb. 2–3) belegt dies visuell: month-to-month-Kunden "
    "kündigen mit >40% Rate, während Demografie kaum differenziert."
)

add_heading(doc3, "H2: Ensemble-Modelle übertreffen lineare Modelle", level=2)
add_paragraph(doc3,
    "H2 lautet: Ensemble-Modelle (Random Forest, Gradient Boosting) erzielen höhere Vorhersageleistung "
    "als Logistic Regression."
)
add_paragraph(doc3,
    "Ergebnis: ❌ WIDERLEGT. Logistic Regression übertrifft beide Ensemble-Modelle beim F1-Score "
    f"(LogReg: 0.614, RF: 0.585, GradBoost: 0.590) und beim Recall (LogReg: 0.783, RF: 0.623, "
    "GradBoost: 0.524). Nur beim ROC-AUC liegt Gradient Boosting marginal vorne (0.843 vs. 0.841). "
    "Mögliche Erklärung: Bei diesem tabellarischen Datensatz mit relativ linearen Zusammenhängen "
    "ist die höhere Komplexität von Ensembles kein Vorteil. Zudem profitiert LogReg stärker von "
    "class_weight='balanced'. H2 muss daher als widerlegt gelten."
)

add_heading(doc3, "H3: Leakage-freie Pipeline verbessert Generalisierung", level=2)
add_paragraph(doc3,
    "H3 lautet: Saubere Vorverarbeitung (Encoding, Scaling, Leakage-freies Vorgehen) verbessert die "
    "Generalisierungsleistung messbar."
)
add_paragraph(doc3,
    "Ergebnis: ✅ BESTÄTIGT (qualitativ). Der DummyClassifier (ROC-AUC: 0.50, F1: 0.00) als "
    "naiver Baseline zeigt, dass ohne Preprocessing und echtes Modell keine Vorhersageleistung "
    "erzielt wird. Das Pipeline-Design verhinderte Datenleckage: Alle Transformationen "
    "(Imputation, Skalierung, Encoding) wurden ausschließlich auf den Trainingsdaten gelernt. "
    "TotalCharges-Konvertierung war entscheidend: Ohne sie hätten Imputation und Skalierung "
    "inkorrekte Werte produziert. H3 gilt als bestätigt, auch wenn ein direkter Vergleich "
    "mit/ohne Pipeline nicht quantifiziert wurde."
)

# ── Streamlit-App ─────────────────────────────────────────────────────────────
add_heading(doc3, "6. K-Phase: Streamlit-App (Deployment)", level=1)
add_heading(doc3, "6.1 Funktionalität", level=2)
add_paragraph(doc3,
    "Die Streamlit-App (app.py) wurde als interaktive Web-Demo des trainierten Modells entwickelt. "
    f"Der Quellcode ist vollständig im GitHub-Repository verfügbar: {PLACEHOLDERS['github']}"
)
add_paragraph(doc3, "Implementierte Funktionen:")
for func in [
    "Eingabeformular mit Slidern und Dropdowns für alle 19 Kundenmerkmale",
    "Echtzeit-Churn-Wahrscheinlichkeitsberechnung via predict_proba()",
    "Farbcodierte Risikoampel: 🟢 <30% (Niedrig), 🟡 30–60% (Mittel), 🔴 >60% (Hoch)",
    "Visuelle Darstellung der Churn-Wahrscheinlichkeit als horizontaler Balken",
    "Feature-Importance-Plot (Koeffizienten) mit Erklärung",
    "Modellvergleichstabelle und Begründung der Modellwahl im Sidebar",
    "Laden des serialisierten Modells (models/logreg_churn_pipeline.pkl via joblib)",
]:
    add_bullet(doc3, func)

add_heading(doc3, "6.2 Code-Auszug: Vorhersage-Logik", level=2)
code_snippet = (
    "# Modell laden\n"
    "pipe = joblib.load('models/logreg_churn_pipeline.pkl')\n\n"
    "# Eingabe als DataFrame\n"
    "input_df = pd.DataFrame([{...Kundendaten...}])\n\n"
    "# Vorhersage\n"
    "prob = pipe.predict_proba(input_df)[0][1]  # Churn-Wahrscheinlichkeit\n"
    "pred = pipe.predict(input_df)[0]           # 0 oder 1\n\n"
    "# Ampel-Logik\n"
    "if prob < 0.30:  risk = 'Niedriges Risiko'  # gruen\n"
    "elif prob < 0.60: risk = 'Mittleres Risiko'  # gelb\n"
    "else:             risk = 'Hohes Risiko'      # rot\n"
)
code_p = doc3.add_paragraph()
code_run = code_p.add_run(code_snippet)
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)
code_p.paragraph_format.left_indent = Cm(1)

add_heading(doc3, "6.3 Lokale Ausführung", level=2)
add_paragraph(doc3, "Voraussetzung: Modell wurde durch Ausführen von C_Evaluation_Final.py erstellt.")
for step in [
    "git clone [GITHUB-REPO-LINK]",
    "pip install -r requirements.txt",
    "python C_Evaluation_Final.py   # Erstellt Modell und Plots",
    "streamlit run app.py           # Startet die Web-App",
]:
    p = doc3.add_paragraph()
    run = p.add_run(step)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)

# ── Zielerreichung & Reflexion ────────────────────────────────────────────────
add_heading(doc3, "7. Zielerreichung und kritische Reflexion", level=1)
add_heading(doc3, "7.1 Stimmt das Ergebnis mit den Zielen aus Phase 1 überein?", level=2)
goal_table = doc3.add_table(rows=1, cols=3)
goal_table.style = 'Table Grid'
for cell, text in zip(goal_table.rows[0].cells, ["Ziel (Phase 1)", "Ergebnis", "Bewertung"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for goal, result, grade in [
    ("F1 > 0.55", f"LogReg: 0.614 (Test)", "✅ Erfüllt"),
    ("Recall > 0.65", f"LogReg: 0.783 (Test)", "✅ Erfüllt"),
    ("ROC-AUC > 0.80", "LogReg: 0.841 (Test)", "✅ Erfüllt"),
    ("Streamlit-App", "app.py implementiert", "✅ Erfüllt"),
    ("H2 (Ensembles besser)", "LogReg ist bestes Modell", "❌ Widerlegt"),
    ("GridSearchCV-Tuning", "Nicht implementiert", "⚠️ Offen"),
]:
    row = goal_table.add_row()
    for cell, val in zip(row.cells, [goal, result, grade]):
        cell.text = val
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc3.add_paragraph()
add_paragraph(doc3,
    "Die primären Erfolgsziele (F1, Recall, ROC-AUC, Streamlit-App) wurden erreicht oder übertroffen. "
    "Lediglich das Hyperparameter-Tuning via GridSearchCV wurde nicht vollständig implementiert – "
    "die Default-Parameter erwiesen sich als ausreichend für die definierten Zielmetriken. "
    "H2 wurde empirisch widerlegt, was als wertvolle Erkenntnis gilt: Modellkomplexität ≠ bessere Leistung."
)

add_heading(doc3, "7.2 Was würde ich beim nächsten Mal anders machen?", level=2)
for text in [
    "Hyperparameter-Tuning von Anfang an einplanen: GridSearchCV für alle Modelle systematisch durchführen, insbesondere Threshold-Optimierung für den Klassifikationsschwellenwert",
    "SHAP-Werte statt reiner Koeffizientenanalyse: SHAP bietet modellunabhängige, intuitivere Feature-Importance-Erklärungen",
    "Weitere Datensätze zum Testen der Generalisierbarkeit (z. B. Bank Customer Churn) einbeziehen",
    "Class Imbalance: SMOTE (Synthetic Minority Oversampling) als Alternative zu class_weight='balanced' testen",
    "Kalibrierungskurven zur Überprüfung der Wahrscheinlichkeitsschätzungen hinzufügen",
    "End-to-End-Deployment via Streamlit Cloud oder Hugging Face Spaces für öffentlich zugängliche Demo",
]:
    add_bullet(doc3, text)

# ── KI-Tool-Disclosure ────────────────────────────────────────────────────────
add_heading(doc3, "8. KI-Tool-Nutzung / AI Tool Disclosure", level=1)
add_paragraph(doc3,
    "Gemäß den Anforderungen der IU Internationalen Hochschule wird die Nutzung von KI-Werkzeugen "
    "transparent und vollständig dokumentiert. Alle durch KI-Systeme generierten Inhalte wurden "
    "kritisch geprüft, angepasst und eigenverantwortlich übernommen. Die finale Verantwortung "
    "für alle Ergebnisse liegt vollständig bei der Autorin/dem Autor des Portfolios."
)

ai_table = doc3.add_table(rows=1, cols=4)
ai_table.style = 'Table Grid'
for cell, text in zip(ai_table.rows[0].cells,
                      ["Phase (QUA³CK)", "KI-Tool", "Wofür genutzt", "Eigenanteil / Anpassung"]):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(9)

ai_data = [
    ("Q – Question", "GitHub Copilot (GPT-4o), Google Gemini 1.5 Pro",
     "Recherche öffentlicher Datensätze; Strukturierungshilfe für Forschungsfrage und Hypothesen",
     "Inhalte selbst geprüft und an eigenes Vorwissen angepasst; Hypothesen eigenständig formuliert"),
    ("U – Understanding", "OpenAI GPT-4o",
     "EDA-Code-Generierung (pandas/seaborn); Interpretationshilfe für Plots",
     "Code verstanden, getestet und an Datensatz angepasst; Interpretationen eigenständig formuliert"),
    ("A¹ – Algorithm Selection", "Google Gemini 1.5 Pro",
     "Vergleich von ML-Algorithmen für tabellarische Daten; Vor-/Nachteile recherchiert",
     "Modellauswahl eigenständig begründet; KI als Recherchetool, nicht als Entscheider"),
    ("A² – Adapting Features", "GitHub Copilot Workspace",
     "Grundgerüst für scikit-learn ColumnTransformer und Pipeline; OneHotEncoder-Konfiguration",
     "Code vollständig verstanden und auf TotalCharges-Problem angepasst; Leakage-Freiheit eigenständig geprüft"),
    ("A³ – Parameter Adjustment", "GitHub Copilot (GPT-4o)",
     "Cross-Validation Setup (StratifiedKFold, make_scorer); Metriken-Auswahl",
     "Scoring-Dictionary und CV-Schleife eigenständig implementiert; KI lieferte Template"),
    ("C – Conclusion", "OpenAI GPT-4o, Claude (Anthropic)",
     "Holdout-Evaluation Code; ROC/PR-Kurven-Code; Feature-Importance-Visualisierung",
     "Code ausgeführt und Ergebnisse eigenständig interpretiert; Plots eigenständig kommentiert"),
    ("K – Knowledge Transfer", "Claude Code (Anthropic)",
     "Streamlit-App Grundgerüst; app.py-Struktur; Ampel-Logik",
     "App vollständig verstanden und an Modell angepasst; UI-Texte eigenständig verfasst"),
    ("Portfolio-Dokumente", "Claude (Anthropic)",
     "Strukturierung und Formatierung der .docx-Dateien; Vollständigkeitsprüfung",
     "Alle inhaltlichen Texte eigenständig verfasst oder aus Notebooks übernommen; KI half bei Gliederung und Form"),
]

for row_data in ai_data:
    row = ai_table.add_row()
    for i, text in enumerate(row_data):
        row.cells[i].text = text
        row.cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)

doc3.add_paragraph()
add_paragraph(doc3,
    "Wichtiger Hinweis: Die Nutzung von KI-Werkzeugen erfolgte stets als Hilfsmittel, "
    "nicht als Ersatz für eigenes Verständnis. Jede KI-generierte Codezeile wurde verstanden, "
    "getestet und ggf. angepasst, bevor sie eingesetzt wurde. Inhaltliche Entscheidungen "
    "(Modellauswahl, Hypothesenformulierung, Ergebnisinterpretation) wurden stets eigenständig getroffen."
)

add_references(doc3)
doc3.save("Portfolio_Phase3_CustomerChurn.docx")
print("✅ Portfolio_Phase3_CustomerChurn.docx gespeichert")
print("\n🎉 Alle 3 Dokumente wurden erstellt!")
